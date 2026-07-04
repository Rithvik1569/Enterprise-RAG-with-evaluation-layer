import os
import logging
from datetime import datetime, timezone
from fastapi import UploadFile, HTTPException, status
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LC_Document

from app.config.settings import settings
from app.models.document import DocumentInDB
from app.utils.file_store import FileStore
from app.utils.vector_store import VectorStoreManager

logger = logging.getLogger("rag_pipeline")


class DocumentProcessor:
    """Orchestrates the document ingestion pipeline:
    Save file -> Extract text -> Clean -> Chunk -> Embed & Index -> Update Metadata DB
    """

    def __init__(self, db):
        self.db = db
        self.file_store = FileStore()
        self.vector_manager = VectorStoreManager()

    async def upload_and_process(
        self,
        file: UploadFile,
        chunk_size: int = settings.CHUNK_SIZE,
        chunk_overlap: int = settings.CHUNK_OVERLAP,
    ):
        """Validates, stores, and indexes a newly uploaded document."""
        filename = file.filename
        file_type = os.path.splitext(filename)[1].lower().strip(".")
        if file_type not in ["pdf", "docx", "txt"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, TXT.",
            )

        # 1. Save as temporary file to calculate checksum and prevent duplicates
        temp_filename = f"temp_{filename}"
        temp_path = await self.file_store.save_file(file, temp_filename)
        checksum = self.file_store.compute_checksum(temp_path)

        # 2. Check duplicate uploads via checksum
        existing_doc = await self.db["documents"].find_one({"checksum": checksum})
        if existing_doc:
            self.file_store.delete_file(temp_path)
            # Return or raise duplicate alert
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"Duplicate upload: '{filename}' has the exact same content as already existing document '{existing_doc['filename']}' (ID: {existing_doc['_id']}).",
            )

        # 3. Rename to permanent filepath
        perm_filename = f"{checksum}_{filename}"
        perm_path = self.file_store.get_file_path(perm_filename)
        
        # If perm file already exists for some reason, delete temp and use perm
        if os.path.exists(perm_path):
            self.file_store.delete_file(temp_path)
        else:
            os.rename(temp_path, perm_path)

        file_size = os.path.getsize(perm_path)

        # 4. Insert metadata record with 'processing' status
        db_doc = DocumentInDB(
            filename=filename,
            file_path=perm_path,
            file_type=file_type,
            size=file_size,
            checksum=checksum,
            processing_status="processing",
            chunk_count=0,
        )
        doc_dict = db_doc.model_dump(by_alias=True)
        await self.db["documents"].insert_one(doc_dict)

        # 5. Run the processing pipeline
        try:
            await self._run_pipeline(db_doc, chunk_size, chunk_overlap)
        except Exception as e:
            logger.exception("Ingestion pipeline failed for document: %s", db_doc.id)
            await self.db["documents"].update_one(
                {"_id": db_doc.id},
                {"$set": {
                    "processing_status": "failed",
                    "error_message": str(e),
                    "updated_at": datetime.now(timezone.utc)
                }}
            )
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Ingestion pipeline failed: {str(e)}",
            )

        # return latest
        final_doc = await self.db["documents"].find_one({"_id": db_doc.id})
        return DocumentInDB(**final_doc)

    async def reprocess_document(
        self,
        doc_id: str,
        chunk_size: int,
        chunk_overlap: int,
    ):
        """Reprocesses an existing document with new chunking configurations."""
        doc_dict = await self.db["documents"].find_one({"_id": doc_id})
        if not doc_dict:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document metadata record not found.",
            )

        if not os.path.exists(doc_dict["file_path"]):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Raw file has been deleted from disk. Cannot reprocess.",
            )

        db_doc = DocumentInDB(**doc_dict)

        # Clean up existing vector store chunks first
        self.vector_manager.delete_document_chunks(db_doc.id)

        # Reset metadata status
        db_doc.processing_status = "processing"
        db_doc.error_message = None
        db_doc.chunk_count = 0
        db_doc.updated_at = datetime.now(timezone.utc)
        
        await self.db["documents"].update_one(
            {"_id": db_doc.id},
            {"$set": {
                "processing_status": "processing",
                "error_message": None,
                "chunk_count": 0,
                "updated_at": db_doc.updated_at
            }}
        )

        # Run pipeline
        try:
            await self._run_pipeline(db_doc, chunk_size, chunk_overlap)
        except Exception as e:
            logger.exception("Reprocessing pipeline failed for document: %s", db_doc.id)
            await self.db["documents"].update_one(
                {"_id": db_doc.id},
                {"$set": {
                    "processing_status": "failed",
                    "error_message": str(e),
                    "updated_at": datetime.now(timezone.utc)
                }}
            )
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Reprocessing failed: {str(e)}",
            )

        final_doc = await self.db["documents"].find_one({"_id": doc_id})
        return DocumentInDB(**final_doc)

    async def _run_pipeline(self, db_doc: DocumentInDB, chunk_size: int, chunk_overlap: int):
        """Internal helper to execute text extraction, cleaning, chunking, and embedding."""
        # A. Extract text
        raw_text = TextExtractor.extract(db_doc.file_path, db_doc.file_type)

        # B. Clean text
        cleaned_text = TextCleaner.clean(raw_text)

        # C. Chunk text
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        chunks = splitter.split_text(cleaned_text)

        if not chunks:
            raise ValueError("Document yielded no text chunks after splitting.")

        # D. Add to vector store
        lc_docs = []
        for idx, chunk in enumerate(chunks):
            lc_docs.append(
                LC_Document(
                    page_content=chunk,
                    metadata={
                        "document_id": db_doc.id,
                        "filename": db_doc.filename,
                        "chunk_index": idx,
                    },
                )
            )

        vectorstore = self.vector_manager.get_vectorstore()
        vectorstore.add_documents(lc_docs)

        # E. Update record status
        await self.db["documents"].update_one(
            {"_id": db_doc.id},
            {"$set": {
                "processing_status": "completed",
                "chunk_count": len(chunks),
                "updated_at": datetime.now(timezone.utc)
            }}
        )
import logging
from app.config.settings import settings
from app.models.document import RetrievalChunk
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
import google.api_core.exceptions

logger = logging.getLogger("rag_pipeline")


class LLMService:
    """Service to interact with Gemini LLM."""

    def __init__(self):
        self.gemini_key = settings.GEMINI_API_KEY
        self.openai_key = None # Removed
        self.groq_key = None # Removed

        if self.gemini_key:
            self.llm = ChatGoogleGenerativeAI(
                model="gemini-2.5-flash",
                google_api_key=self.gemini_key,
                temperature=0.2,
                max_tokens=800
            )
        else:
            self.llm = None

    def build_prompt(self, query: str, chunks: list[RetrievalChunk]) -> str:
        """Constructs a context-aware prompt using the retrieved document chunks."""
        if not chunks:
            return f"User Question: {query}\n\nNo document context was found in the database. Please answer to the best of your knowledge but specify that no relevant documents were found."

        context_str = ""
        for i, chunk in enumerate(chunks, 1):
            context_str += f"--- Source {i}: {chunk.filename} (Chunk {chunk.chunk_index}) ---\n"
            context_str += f"{chunk.text}\n\n"

        prompt = (
            f"=== DOCUMENT CONTEXT ===\n{context_str}========================\n\n"
            f"User Question: {query}\n\n"
            "Answer:"
        )
        return prompt

    async def generate_answer(self, query: str, chunks: list[RetrievalChunk]) -> str:
        """Generates an answer from Gemini."""
        prompt = self.build_prompt(query, chunks)
        
        if not self.llm:
            logger.warning("No functional LLM API key available or calls failed. Using fallback mock responder.")
            return self._generate_mock_response(query, chunks, error_msg="GEMINI_API_KEY is not set.")

        try:
            logger.info("Generating response using Gemini API...")
            messages = [
                SystemMessage(content="You are a helpful and precise RAG assistant. First, try to answer the user's question using the provided document context below. If the context contains the answer, refer to the sources where appropriate (e.g., [Source 1], [Source 2]). If the context does not contain the answer, answer to the best of your knowledge like ChatGPT, but clearly state that your answer is based on general knowledge and not found in the provided documents."),
                HumanMessage(content=prompt)
            ]
            response = await self.llm.ainvoke(messages)
            return response.content
        except google.api_core.exceptions.GoogleAPICallError as e:
            logger.error("Failed to generate response with Gemini API: %s. Falling back...", str(e))
            return self._generate_mock_response(query, chunks, error_msg=str(e))
        except Exception as e:
            logger.error("Unknown error in LLM generation: %s", str(e))
            return self._generate_mock_response(query, chunks, error_msg=str(e))

    async def generate_text(self, prompt: str) -> str:
        """Generates content from Gemini for any arbitrary prompt."""
        if not self.llm:
            return "mock_eval_response"
            
        try:
            logger.info("Generating text using Gemini API...")
            response = await self.llm.ainvoke([HumanMessage(content=prompt)])
            return response.content
        except Exception as e:
            logger.error("Failed to generate text with Gemini API: %s. Falling back...", str(e))
            return "mock_eval_response"

    def _generate_mock_response(self, query: str, chunks: list[RetrievalChunk], error_msg: str = None) -> str:
        """Generates a mock response simulating LLM behavior based on retrieved chunks."""
        if not chunks:
            return (
                "ðŸ‘‹ Hello! I am running in development fallback mode. I searched the database but found no relevant document chunks. "
                "Please upload a document to the Knowledge Base first so I can find matching information to answer your question!"
            )

        citations_list = [f"**{c.filename}** (Chunk {c.chunk_index}, relevance score: {c.score:.4f})" for c in chunks]
        
        matched_sentences = []
        query_words = [w.lower() for w in query.split() if len(w) > 3]
        
        for c in chunks:
            for line in c.text.split('.'):
                line = line.strip()
                if not line:
                    continue
                if any(qw in line.lower() for qw in query_words):
                    matched_sentences.append(line)
                    if len(matched_sentences) >= 3:
                        break
            if len(matched_sentences) >= 3:
                break

        if not matched_sentences:
            first_chunk_text = chunks[0].text.replace('\n', ' ')
            matched_sentences = first_chunk_text.split('.')[:2]

        summary = ". ".join([s.strip() for s in matched_sentences if s.strip()]) + "."
        
        error_note = f"\n\n*API Error Details:* `{error_msg}`" if error_msg else ""
        
        answer = (
            "âš™ï¸ **[Development Fallback Mode]**\n\n"
            f"Based on the retrieved document context, here is what I found:\n\n"
            f"> \"...{summary}...\"\n\n"
            "This information was retrieved from the following sources:\n" + 
            "\n".join([f"- [Source {i+1}] {cite}" for i, cite in enumerate(citations_list)]) + "\n\n"
            "*Note: The Gemini API is currently unavailable, so I cannot provide a proper answer.*" + error_note + "\n"
            "*Please check your `.env` API keys and Quota limits!*"
        )
        return answer
import logging
from app.utils.vector_store import VectorStoreManager
from app.models.document import RetrievalChunk

logger = logging.getLogger("rag_pipeline")


class RetrievalService:
    """Handles vector search queries and formats matched chunks with relevance scores."""

    def __init__(self):
        self.vector_manager = VectorStoreManager()

    async def retrieve(
        self,
        query: str,
        top_k: int = 4,
        document_id: str | None = None,
    ) -> list[RetrievalChunk]:
        """Queries ChromaDB vector index and returns similarity matched chunks."""
        vectorstore = self.vector_manager.get_vectorstore()
        
        # Define filter if document_id is provided
        search_filter = None
        if document_id:
            search_filter = {"document_id": document_id}

        logger.info("Executing retrieval search: query='%s', top_k=%d, filter=%s", query, top_k, search_filter)

        # Retrieve documents with relevance scores
        results = vectorstore.similarity_search_with_score(
            query,
            k=top_k,
            filter=search_filter,
        )

        retrieved_chunks = []
        for doc, score in results:
            retrieved_chunks.append(
                RetrievalChunk(
                    text=doc.page_content,
                    document_id=doc.metadata.get("document_id", ""),
                    filename=doc.metadata.get("filename", ""),
                    chunk_index=int(doc.metadata.get("chunk_index", 0)),
                    score=float(score),
                )
            )

        return retrieved_chunks
import re


class TextCleaner:
    """Utility service to clean and normalize raw extracted document text."""

    @staticmethod
    def clean(text: str) -> str:
        """Cleans and normalizes raw text extracted from documents.
        
        - Standardizes newlines to '\n'
        - Replaces multiple horizontal spaces/tabs with a single space
        - Collapses 3 or more consecutive newlines to a double newline
        - Strips whitespace from lines
        """
        if not text:
            return ""

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Replace multiple consecutive spaces/tabs within a line with a single space
        text = re.sub(r"[ \t]+", " ", text)

        # Strip spaces from each line, but preserve newlines
        lines = [line.strip() for line in text.split("\n")]
        
        # Re-join lines
        text = "\n".join(lines)

        # Compress 3 or more consecutive newlines into a maximum of 2 (a blank line)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()
import os
import pypdf
import docx
import logging

logger = logging.getLogger("rag_pipeline")


class TextExtractor:
    """Extracts raw text from PDF, DOCX, and TXT files."""

    @staticmethod
    def extract(file_path: str, file_type: str) -> str:
        """Extracts text based on the file type."""
        file_type = file_type.lower().strip(".")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")

        logger.info("Extracting text from %s (type: %s)", file_path, file_type)

        if file_type == "pdf":
            return TextExtractor._extract_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return TextExtractor._extract_docx(file_path)
        elif file_type == "txt":
            return TextExtractor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        extracted_text = []
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                num_pages = len(reader.pages)
                logger.info("PDF has %d pages", num_pages)
                
                for idx, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text.append(page_text)
                    else:
                        logger.debug("Page %d of PDF returned no text", idx + 1)
        except Exception as e:
            logger.error("Error reading PDF %s: %s", file_path, str(e))
            raise ValueError(f"Corrupted or invalid PDF file: {str(e)}")

        text = "\n\n".join(extracted_text)
        if not text.strip():
            raise ValueError("Extracted text from PDF is empty. File might contain only scanned images or is empty.")
        return text

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            # Check tables too
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_text.append(" | ".join(row_data))
            
            if table_text:
                text += "\n\n=== Table Data ===\n" + "\n".join(table_text)
                
        except Exception as e:
            logger.error("Error reading DOCX %s: %s", file_path, str(e))
            raise ValueError(f"Corrupted or invalid DOCX file: {str(e)}")

        if not text.strip():
            raise ValueError("Extracted text from DOCX is empty.")
        return text

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                    logger.info("Successfully read TXT file using %s encoding", enc)
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error("Error reading TXT %s with encoding %s: %s", file_path, enc, str(e))
                raise ValueError(f"Error reading TXT file: {str(e)}")

        raise ValueError("Could not decode TXT file with any supported encoding (tried UTF-8, Latin-1, CP1252).")
