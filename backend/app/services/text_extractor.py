import os
import pypdf
import docx
import logging

logger = logging.getLogger("rag_pipeline")


class TextExtractor:
    """Extracts raw text from PDF, DOCX, and TXT files."""

    @staticmethod
    def extract(file_path: str, file_type: str) -> str:
        """Extracts text based on the file type."""
        file_type = file_type.lower().strip(".")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")

        logger.info("Extracting text from %s (type: %s)", file_path, file_type)

        if file_type == "pdf":
            return TextExtractor._extract_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return TextExtractor._extract_docx(file_path)
        elif file_type == "txt":
            return TextExtractor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        extracted_text = []
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                num_pages = len(reader.pages)
                logger.info("PDF has %d pages", num_pages)
                
                for idx, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text.append(page_text)
                    else:
                        logger.debug("Page %d of PDF returned no text", idx + 1)
        except Exception as e:
            logger.error("Error reading PDF %s: %s", file_path, str(e))
            raise ValueError(f"Corrupted or invalid PDF file: {str(e)}")

        text = "\n\n".join(extracted_text)
        if not text.strip():
            raise ValueError("Extracted text from PDF is empty. File might contain only scanned images or is empty.")
        return text

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            # Check tables too
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_text.append(" | ".join(row_data))
            
            if table_text:
                text += "\n\n=== Table Data ===\n" + "\n".join(table_text)
                
        except Exception as e:
            logger.error("Error reading DOCX %s: %s", file_path, str(e))
            raise ValueError(f"Corrupted or invalid DOCX file: {str(e)}")

        if not text.strip():
            raise ValueError("Extracted text from DOCX is empty.")
        return text

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                    logger.info("Successfully read TXT file using %s encoding", enc)
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error("Error reading TXT %s with encoding %s: %s", file_path, enc, str(e))
                raise ValueError(f"Error reading TXT file: {str(e)}")

        raise ValueError("Could not decode TXT file with any supported encoding (tried UTF-8, Latin-1, CP1252).")
